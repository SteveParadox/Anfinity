import io
import sys
import unittest
from pathlib import Path
from uuid import uuid4

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.ingestion.chunker import Chunker
from app.ingestion.parsers.html import HTMLParser
from app.ingestion.parsers.pdf import PDFParser
from app.ingestion.parsers.text import TextParser
from app.ingestion.parsers.word import WordParser
from app.ingestion.source_locations import enrich_citation_metadata
from app.services.answer_generator import AnswerGenerator, RetrievedChunk as AnswerChunk
from app.services.top_k_retriever import TopKRetriever


class SourceLocationPipelineTests(unittest.TestCase):
    def setUp(self) -> None:
        self.chunker = Chunker(chunk_size=400, chunk_overlap=40)

    def test_markdown_parser_and_chunker_capture_lines_headings_and_char_offsets(self):
        markdown = (
            "# Dyslexia Overview\n"
            "\n"
            "Dyslexia affects reading fluency and decoding.\n"
            "Structured literacy interventions can help.\n"
            "\n"
            "## Support\n"
            "\n"
            "Accommodations may include extra time and guided reading practice.\n"
        ).encode("utf-8")

        parsed = TextParser().parse(markdown)
        chunks = self.chunker.chunk_parsed_document(
            parsed,
            metadata={
                "source_file_name": "notes.md",
                "document_title": "notes.md",
                "source_type": "upload",
            },
        )

        self.assertTrue(parsed.segments)
        self.assertGreaterEqual(len(chunks), 1)
        metadata = chunks[0].metadata
        self.assertEqual(metadata["source_file_name"], "notes.md")
        self.assertEqual(metadata["heading_path_start"], ["Dyslexia Overview"])
        self.assertEqual(metadata["section_title_start"], "Dyslexia Overview")
        self.assertEqual(metadata["line_start"], 1)
        self.assertGreaterEqual(metadata["line_end"], 4)
        self.assertIn("citation_label", metadata)
        self.assertIn("notes.md", metadata["citation_label"])
        self.assertIn("heading 'Dyslexia Overview'", metadata["citation_label"])
        self.assertIn("char_start", metadata)
        self.assertIn("char_end", metadata)

    def test_docx_parser_captures_heading_hierarchy_and_paragraph_indexes(self):
        document = Document()
        document.add_heading("Reading Disorders", level=1)
        document.add_paragraph("Dyslexia is commonly discussed in literacy interventions.")
        document.add_heading("Classroom Support", level=2)
        document.add_paragraph("Teachers may provide structured reading support.")

        buffer = io.BytesIO()
        document.save(buffer)

        parsed = WordParser().parse(buffer.getvalue())
        chunks = self.chunker.chunk_parsed_document(
            parsed,
            metadata={
                "source_file_name": "Medical Learning Disorders.docx",
                "document_title": "Medical Learning Disorders.docx",
                "source_type": "upload",
            },
        )

        metadata = chunks[0].metadata
        self.assertEqual(metadata["heading_path_start"], ["Reading Disorders"])
        self.assertEqual(metadata["section_title_start"], "Reading Disorders")
        self.assertEqual(metadata["paragraph_start"], 1)
        self.assertIsNone(metadata.get("page_number"))

    def test_html_parser_captures_dom_context_and_heading_path(self):
        html = b"""
        <html>
          <body>
            <article id="article">
              <h1>Dyslexia Overview</h1>
              <section id="support">
                <h2>Classroom Support</h2>
                <p>Guided reading and structured literacy are common supports.</p>
              </section>
            </article>
          </body>
        </html>
        """

        parsed = HTMLParser().parse(html)
        chunks = self.chunker.chunk_parsed_document(
            parsed,
            metadata={
                "source_file_name": "dyslexia.html",
                "document_title": "dyslexia.html",
                "source_type": "upload",
            },
        )

        self.assertTrue(parsed.segments)
        metadata = chunks[0].metadata
        self.assertEqual(metadata["heading_path"], ["Dyslexia Overview"])
        self.assertEqual(metadata["heading_path_end"], ["Dyslexia Overview", "Classroom Support"])
        self.assertTrue(metadata["spans_multiple_sections"])
        self.assertIn("article#article", metadata["dom_path_start"])
        self.assertEqual(metadata["paragraph_start"], 1)
        self.assertIn("citation_label", metadata)

    def test_pdf_chunk_can_span_pages_without_faking_single_page_number(self):
        pdf = fitz.open()
        page_one = pdf.new_page()
        page_one.insert_text((72, 72), "Reading Disorders", fontsize=20)
        page_one.insert_text(
            (72, 110),
            "Dyslexia is a learning disorder that affects decoding and reading fluency.",
            fontsize=11,
        )
        page_two = pdf.new_page()
        page_two.insert_text(
            (72, 72),
            "Interventions often focus on structured literacy and phonics practice.",
            fontsize=11,
        )
        pdf_bytes = pdf.tobytes()
        pdf.close()

        parsed = PDFParser().parse(pdf_bytes)
        chunks = self.chunker.chunk_parsed_document(
            parsed,
            metadata={
                "source_file_name": "Medical Learning Disorders.pdf",
                "document_title": "Medical Learning Disorders.pdf",
                "source_type": "upload",
            },
        )

        self.assertEqual(len(chunks), 1)
        metadata = chunks[0].metadata
        self.assertEqual(metadata["page_start"], 1)
        self.assertEqual(metadata["page_end"], 2)
        self.assertNotIn("page_number", metadata)
        self.assertTrue(metadata["spans_multiple_pages"])
        self.assertIn("pages 1-2", metadata["citation_label"])

    def test_retrieval_returns_citation_rich_chunk_metadata(self):
        class FakeEmbeddingService:
            def embed_query(self, query):
                return [0.1, 0.2, 0.3]

            def get_dimension(self):
                return 3

        class FakeVectorDB:
            def resolve_collection_name(self, base_name, embedding_dim=None):
                return str(base_name)

            def search_similar(self, **kwargs):
                return [
                    {
                        "id": "chunk-1",
                        "similarity": 0.91,
                        "payload": {
                            "chunk_id": "chunk-1",
                            "document_id": "doc-1",
                            "document_title": "Medical Learning Disorders.pdf",
                            "chunk_text": "This medical book discusses dyslexia and reading disorders.",
                            "source_type": "upload",
                            "chunk_index": 0,
                            "token_count": 12,
                            "metadata": {
                                "source_file_name": "Medical Learning Disorders.pdf",
                                "page_number": 42,
                                "section_title": "Reading Disorders",
                            },
                        },
                    }
                ]

        retriever = object.__new__(TopKRetriever)
        retriever.db = None
        retriever.top_k = 3
        retriever.similarity_threshold = 0.0
        retriever.embedding_service = FakeEmbeddingService()
        retriever.vector_db = FakeVectorDB()

        result = retriever.retrieve(
            query="medical book on dyslexia",
            workspace_id=uuid4(),
            top_k=1,
            similarity_threshold=0.0,
        )

        self.assertEqual(len(result.chunks), 1)
        metadata = result.chunks[0].metadata
        self.assertEqual(metadata["page_number"], 42)
        self.assertEqual(metadata["section_title"], "Reading Disorders")
        self.assertIn("page 42", metadata["citation_label"])
        self.assertEqual(metadata["source_location"]["page_number"], 42)

    def test_answer_context_uses_real_citation_metadata_and_missing_fields_degrade_honestly(self):
        generator = AnswerGenerator()
        chunk = AnswerChunk(
            chunk_id="chunk-1",
            document_id="doc-1",
            similarity=0.88,
            text="Dyslexia interventions often emphasize structured literacy.",
            source_type="document",
            chunk_index=0,
            document_title="Medical Learning Disorders.pdf",
            token_count=9,
            metadata=enrich_citation_metadata(
                {
                    "source_file_name": "Medical Learning Disorders.pdf",
                    "page_number": 42,
                    "section_title": "Reading Disorders",
                },
                document_title="Medical Learning Disorders.pdf",
                source_type="document",
            ),
        )

        context = generator._build_context([chunk], include_citations=True)
        citations = generator._extract_citations([chunk], "Medical Learning Disorders.pdf discusses dyslexia.")

        self.assertIn("Citation: Medical Learning Disorders.pdf, page 42, section 'Reading Disorders'", context)
        self.assertEqual(citations[0].citation_label, "Medical Learning Disorders.pdf, page 42, section 'Reading Disorders'")
        self.assertEqual(citations[0].source_location["page_number"], 42)

        degraded = enrich_citation_metadata(
            {"source_file_name": "notes.md", "heading_path": ["Dyslexia Overview"]},
            document_title="notes.md",
            source_type="document",
        )
        self.assertEqual(degraded["citation_label"], "notes.md, heading 'Dyslexia Overview'")
        self.assertNotIn("page", degraded["citation_label"])


if __name__ == "__main__":
    unittest.main()
