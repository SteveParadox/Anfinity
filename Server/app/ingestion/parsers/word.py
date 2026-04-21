"""Word document parser with heading and paragraph locations."""

from __future__ import annotations

import io
import re
from typing import List

from docx import Document

from app.ingestion.parsers.base import DocumentParser, ParsedDocument


_HEADING_LEVEL_RE = re.compile(r"heading\s+(\d+)", re.IGNORECASE)


class WordParser(DocumentParser):
    """Word document parser (.docx)."""

    def parse(self, file_bytes: bytes) -> ParsedDocument:
        doc = Document(io.BytesIO(file_bytes))

        metadata = {}
        core_props = doc.core_properties
        metadata["title"] = core_props.title
        metadata["author"] = core_props.author
        metadata["subject"] = core_props.subject
        metadata["created"] = str(core_props.created) if core_props.created else None
        metadata["modified"] = str(core_props.modified) if core_props.modified else None
        metadata["content_format"] = "docx"

        segments = []
        heading_path: List[str] = []
        paragraph_index = 0
        block_index = 0

        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue

            style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
            heading_level = self._heading_level(style_name)

            if heading_level is not None:
                heading_path = heading_path[: heading_level - 1] + [text]
                block_index += 1
                segment = self._build_segment(
                    text,
                    segment_type="heading",
                    source_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    block_index=block_index,
                    heading_path=list(heading_path),
                    heading_path_start=list(heading_path),
                    heading_path_end=list(heading_path),
                    section_title=text,
                    section_title_start=text,
                    section_title_end=text,
                    extraction_confidence=0.95,
                )
            else:
                paragraph_index += 1
                block_index += 1
                segment = self._build_segment(
                    text,
                    segment_type="paragraph",
                    source_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    paragraph_index=paragraph_index,
                    block_index=block_index,
                    heading_path=list(heading_path),
                    heading_path_start=list(heading_path),
                    heading_path_end=list(heading_path),
                    section_title=heading_path[-1] if heading_path else None,
                    section_title_start=heading_path[-1] if heading_path else None,
                    section_title_end=heading_path[-1] if heading_path else None,
                    extraction_confidence=0.95,
                )
            if segment:
                segments.append(segment)

        for table_index, table in enumerate(doc.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not cells:
                    continue
                paragraph_index += 1
                block_index += 1
                table_text = " | ".join(cells)
                segment = self._build_segment(
                    table_text,
                    segment_type="table_row",
                    source_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    paragraph_index=paragraph_index,
                    block_index=block_index,
                    heading_path=list(heading_path),
                    heading_path_start=list(heading_path),
                    heading_path_end=list(heading_path),
                    section_title=heading_path[-1] if heading_path else None,
                    section_title_start=heading_path[-1] if heading_path else None,
                    section_title_end=heading_path[-1] if heading_path else None,
                    extraction_confidence=0.9,
                    metadata={"table_index": table_index, "row_index": row_index},
                )
                if segment:
                    segments.append(segment)

        title = metadata.get("title")
        if not title:
            for segment in segments:
                if segment.segment_type == "heading":
                    title = segment.text
                    break

        return self._build_document_from_segments(
            segments=segments,
            metadata=metadata,
            title=title,
            author=metadata.get("author"),
        )

    def _heading_level(self, style_name: str) -> int | None:
        match = _HEADING_LEVEL_RE.search(style_name or "")
        if not match:
            return None
        return max(1, int(match.group(1)))
